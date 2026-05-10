"""
Statistical Regression Analysis App
Simple & Multiple Linear Regression — SPSS-equivalent output

Fixes applied:
  1. KS normality test now uses Lilliefors correction (statsmodels) — matches SPSS exactly
  2. Glejser KeyError fixed — use .iloc[] instead of integer index on params Series
  3. Theme switched to clean light theme for full readability
"""

import streamlit as st
import pandas as pd
import numpy as np
from scipy import stats
from scipy.stats import shapiro, norm
from statsmodels.stats.diagnostic import lilliefors
import statsmodels.api as sm
from statsmodels.stats.stattools import durbin_watson
from statsmodels.stats.outliers_influence import variance_inflation_factor
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from io import BytesIO
import warnings
warnings.filterwarnings("ignore")

st.set_page_config(page_title="Regression Analyzer", page_icon="📊",
                   layout="wide", initial_sidebar_state="collapsed")

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=DM+Serif+Display:ital@0;1&family=DM+Sans:wght@300;400;500;600&family=DM+Mono:wght@400;500&display=swap');
html, body, [class*="css"], .stApp {
    font-family: 'DM Sans', sans-serif !important;
    background-color: #f8f7f4 !important;
    color: #1a1a2e !important;
}
st.markdown('<p class="main-title">Nama Aplikasi Anda</p>', unsafe_allow_html=True)
st.markdown('<p class="main-subtitle">Deskripsi singkat aplikasi</p>', unsafe_allow_html=True)
.main-title { font-family:'DM Serif Display',serif; font-size:2.8rem; color:#1a1a2e; letter-spacing:-1px; line-height:1.1; margin-bottom:0.2rem; }
.main-subtitle { font-weight:400; color:#5a5a7a; font-size:1rem; margin-bottom:1.5rem; }
.step-badge { display:inline-block; background:#e8e2d8; border:1px solid #c9b99a; color:#7a6a50; font-family:'DM Mono',monospace; font-size:0.72rem; padding:0.18rem 0.7rem; border-radius:20px; margin-bottom:0.4rem; letter-spacing:1.5px; text-transform:uppercase; }
.section-header { font-family:'DM Serif Display',serif; font-size:1.5rem; color:#1a1a2e; border-bottom:2px solid #c9b99a; padding-bottom:0.4rem; margin:1.2rem 0 0.8rem 0; }
.interp-box { background:#ffffff; border-left:4px solid #b5965a; padding:1rem 1.3rem; border-radius:0 8px 8px 0; margin:0.8rem 0; font-size:0.93rem; color:#2a2a3e; line-height:1.75; box-shadow:0 1px 4px rgba(0,0,0,0.06); }
.badge-pass { background:#e8f5e9; border:1.5px solid #66bb6a; color:#2e7d32; padding:0.3rem 1rem; border-radius:20px; font-size:0.85rem; font-weight:700; display:inline-block; margin:0.4rem 0; }
.badge-fail { background:#fdecea; border:1.5px solid #ef5350; color:#c62828; padding:0.3rem 1rem; border-radius:20px; font-size:0.85rem; font-weight:700; display:inline-block; margin:0.4rem 0; }
.equation-box { background:#fffdf7; border:1.5px solid #c9b99a; border-radius:10px; padding:1.1rem 1.5rem; font-family:'DM Mono',monospace; font-size:1.15rem; color:#1a1a2e; text-align:center; margin:1rem 0; box-shadow:0 1px 4px rgba(0,0,0,0.05); }
.divider { border:none; border-top:1px solid #ddd8ce; margin:1.5rem 0; }
[data-testid="metric-container"] { background:#ffffff; border:1px solid #e0d8cc; border-radius:10px; padding:0.8rem 1rem; box-shadow:0 1px 3px rgba(0,0,0,0.05); }
[data-testid="metric-container"] label { color:#6a6a8a !important; font-size:0.76rem !important; text-transform:uppercase; letter-spacing:0.8px; }
[data-testid="metric-container"] [data-testid="metric-value"] { font-family:'DM Mono',monospace !important; color:#1a1a2e !important; font-size:1.3rem !important; }
.stRadio label { color:#2a2a3e !important; }
.stSelectbox label, .stMultiSelect label { color:#5a5a7a !important; font-size:0.85rem !important; }
h3 { color:#1a1a2e !important; font-size:1.15rem !important; font-weight:600 !important; }
h4 { color:#2a2a3e !important; }
.stDownloadButton > button { background:#1a1a2e !important; color:#f8f7f4 !important; font-weight:600 !important; border:none !important; border-radius:8px !important; padding:0.55rem 1.5rem !important; }
.stDownloadButton > button:hover { background:#2e2e50 !important; }

a[href*="github.com"] { display: none !important; }
[data-testid="stToolbar"] { display: none !important; }
#MainMenu { visibility: hidden !important; }
.stDeployButton { display: none !important; }
</style>
""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════════════════════
# STATISTICAL FUNCTIONS
# ═══════════════════════════════════════════════════════════════════════════════

def run_normality(residuals):
    """
    FIX 1: KS test now uses statsmodels lilliefors() — matches SPSS Lilliefors-corrected KS.
    Plain scipy kstest() gives different (lower) p-values than SPSS.
    """
    n   = len(residuals)
    res = np.array(residuals, dtype=float)
    sw_stat, sw_p = shapiro(res)
    ks_stat, ks_p = lilliefors(res, dist='norm', pvalmethod='approx')
    primary       = "Shapiro-Wilk" if n <= 50 else "Kolmogorov-Smirnov"
    primary_stat  = sw_stat if n <= 50 else ks_stat
    primary_p     = sw_p    if n <= 50 else ks_p
    return {"n": n,
            "sw_stat": sw_stat, "sw_p": sw_p,
            "ks_stat": ks_stat, "ks_p": ks_p,
            "primary": primary, "primary_stat": primary_stat,
            "primary_p": primary_p, "passed": primary_p > 0.05}


def run_linearity(x_series, y_series):
    x = np.array(x_series, dtype=float)
    y = np.array(y_series, dtype=float)
    n = len(y)
    grand_mean = y.mean()
    x_c  = sm.add_constant(x)
    ols  = sm.OLS(y, x_c).fit()
    ss_reg = float(ols.ess)
    ss_res = float(ols.ssr)
    groups = {}
    for xi, yi in zip(x, y):
        groups.setdefault(round(float(xi), 10), []).append(yi)
    n_groups  = len(groups)
    df_within = n - n_groups
    if n_groups >= n - 1 or df_within <= 0:
        return {"ss_reg": ss_reg, "ss_dev": np.nan, "ss_within": np.nan,
                "df_lin": 1, "df_dev": 0, "df_within": n-2,
                "ms_lin": ss_reg, "ms_dev": np.nan, "ms_within": ss_res/(n-2),
                "F_lin": float(ols.fvalue), "p_lin": float(ols.f_pvalue),
                "F_dev": np.nan, "p_dev": np.nan,
                "passed": float(ols.f_pvalue) < 0.05, "continuous_x": True}
    ss_within  = sum(np.sum((np.array(v)-np.mean(v))**2) for v in groups.values())
    ss_between = sum(len(v)*(np.mean(v)-grand_mean)**2 for v in groups.values())
    df_groups  = n_groups - 1
    ss_dev     = ss_between - ss_reg
    df_lin     = 1
    df_dev     = df_groups - df_lin
    ms_within  = ss_within / df_within if df_within > 0 else np.nan
    ms_lin     = ss_reg / df_lin
    ms_dev     = ss_dev / df_dev if df_dev > 0 else np.nan
    F_lin = ms_lin / ms_within if (ms_within and ms_within > 0) else np.nan
    F_dev = (ms_dev / ms_within
             if (ms_dev is not None and not np.isnan(ms_dev)
                 and ms_within and ms_within > 0 and df_dev > 0) else np.nan)
    p_lin = float(1 - stats.f.cdf(F_lin, df_lin, df_within)) if not np.isnan(F_lin) else np.nan
    p_dev = (float(1 - stats.f.cdf(F_dev, df_dev, df_within))
             if (not np.isnan(F_dev) and df_dev > 0) else np.nan)
    passed = ((not np.isnan(p_lin) and p_lin < 0.05) and
              (np.isnan(p_dev) if np.isnan(p_dev) else p_dev > 0.05))
    ms_between = ss_between / df_groups
    F_between  = ms_between / ms_within if (ms_within and ms_within > 0) else np.nan
    p_between  = float(1 - stats.f.cdf(F_between, df_groups, df_within)) if not np.isnan(F_between) else np.nan
    return {"ss_reg": ss_reg, "ss_dev": ss_dev, "ss_within": ss_within,
            "ss_between": ss_between, "df_groups": df_groups,
            "ms_between": ms_between, "F_between": F_between, "p_between": p_between,
            "df_lin": df_lin, "df_dev": df_dev, "df_within": df_within,
            "ms_lin": ms_lin, "ms_dev": ms_dev, "ms_within": ms_within,
            "F_lin": F_lin, "p_lin": p_lin, "F_dev": F_dev, "p_dev": p_dev,
            "passed": passed, "continuous_x": False}


def run_vif(X_df):
    X = sm.add_constant(X_df.values)
    rows = []
    for i, col in enumerate(X_df.columns):
        vif = float(variance_inflation_factor(X, i + 1))
        rows.append({"Variable": col, "Tolerance": 1.0/vif if vif else np.nan, "VIF": vif})
    return pd.DataFrame(rows)


def run_glejser(residuals, X_df):
    """
    FIX 2: Use .iloc[] for all model.params/bse/tvalues/pvalues access.
    sm.add_constant() on a numpy array produces a named Series index
    ('const','x1','x2',...) in statsmodels — integer key [0] raises KeyError
    in pandas >= 2.0.
    """
    abs_res = np.abs(np.array(residuals, dtype=float))
    X_c     = sm.add_constant(X_df.values)
    model   = sm.OLS(abs_res, X_c).fit()
    p_arr  = np.asarray(model.params)
    b_arr  = np.asarray(model.bse)
    t_arr  = np.asarray(model.tvalues)
    pv_arr = np.asarray(model.pvalues)
    rows = [{"Term": "(Constant)",
             "B":          float(p_arr[0]),
             "Std. Error": float(b_arr[0]),
             "t":          float(t_arr[0]),
             "Sig.":       float(pv_arr[0])}]
    for i, col in enumerate(X_df.columns):
        rows.append({"Term": col,
                     "B":          float(p_arr[i+1]),
                     "Std. Error": float(b_arr[i+1]),
                     "t":          float(t_arr[i+1]),
                     "Sig.":       float(pv_arr[i+1])})
    passed = all(r["Sig."] > 0.05 for r in rows if r["Term"] != "(Constant)")
    return pd.DataFrame(rows), passed, model


def run_regression(y, X_df):
    X_c   = sm.add_constant(X_df.values)
    model = sm.OLS(y.values, X_c).fit()
    n = len(y); k = X_df.shape[1]
    R      = float(np.sqrt(model.rsquared))
    R2     = float(model.rsquared)
    adj_R2 = float(model.rsquared_adj)
    se_est = float(np.sqrt(model.mse_resid))
    dw     = float(durbin_watson(model.resid))
    ss_reg = float(model.ess); ss_res = float(model.ssr); ss_tot = float(model.centered_tss)
    df_reg = k; df_res = n-k-1; df_tot = n-1
    ms_reg = ss_reg/df_reg; ms_res = ss_res/df_res
    F = float(model.fvalue); p_F = float(model.f_pvalue)
    anova = pd.DataFrame([
        {"": "Regression", "Sum of Squares": ss_reg, "df": df_reg, "Mean Square": ms_reg, "F": F,  "Sig.": p_F},
        {"": "Residual",   "Sum of Squares": ss_res, "df": df_res, "Mean Square": ms_res, "F": "",  "Sig.": ""},
        {"": "Total",      "Sum of Squares": ss_tot, "df": df_tot, "Mean Square": "",     "F": "",  "Sig.": ""},
    ])
    y_std = float(y.std(ddof=1))
    # Use .iloc[] throughout — FIX 2 applied here too
    p2  = np.asarray(model.params)
    b2  = np.asarray(model.bse)
    t2  = np.asarray(model.tvalues)
    pv2 = np.asarray(model.pvalues)
    coef_rows = [{"Model": "(Constant)",
                  "B": float(p2[0]), "Std. Error": float(b2[0]),
                  "Beta": "", "t": float(t2[0]), "Sig.": float(pv2[0])}]
    for i, col in enumerate(X_df.columns):
        x_std = float(X_df[col].std(ddof=1))
        beta  = float(p2[i+1]) * (x_std / y_std)
        coef_rows.append({"Model": col,
                          "B": float(p2[i+1]), "Std. Error": float(b2[i+1]),
                          "Beta": beta, "t": float(t2[i+1]),
                          "Sig.": float(pv2[i+1])})
    return {"model": model,
            "R": R, "R2": R2, "adj_R2": adj_R2, "se_est": se_est, "dw": dw,
            "ss_reg": ss_reg, "ss_res": ss_res, "ss_tot": ss_tot,
            "df_reg": df_reg, "df_res": df_res, "df_tot": df_tot,
            "ms_reg": ms_reg, "ms_res": ms_res, "F": F, "p_F": p_F,
            "anova": anova, "coef_df": pd.DataFrame(coef_rows),
            "residuals": pd.Series(model.resid, name="Residuals"),
            "fitted":    pd.Series(model.fittedvalues, name="Fitted"),
            "n": n, "k": k}


# ═══════════════════════════════════════════════════════════════════════════════
# PLOTS  —  light background
# ═══════════════════════════════════════════════════════════════════════════════

BG = "#f8f7f4"; AX = "#ffffff"; AC = "#b5965a"; DK = "#1a1a2e"; MT = "#8a8a9a"

def style_ax(ax, fig):
    fig.patch.set_facecolor(BG); ax.set_facecolor(AX)
    ax.tick_params(colors=MT, labelsize=8.5)
    ax.xaxis.label.set_color(DK); ax.yaxis.label.set_color(DK); ax.title.set_color(DK)
    for sp in ax.spines.values(): sp.set_edgecolor("#d0ccc4")

def plot_normality(residuals):
    fig, axes = plt.subplots(1, 2, figsize=(11, 4.2))
    fig.patch.set_facecolor(BG)
    ax1 = axes[0]; ax1.set_facecolor(AX)
    nb = min(20, max(5, len(residuals)//3))
    ax1.hist(residuals, bins=nb, color=AC, alpha=0.75, edgecolor="white", linewidth=0.6)
    mu, sigma = float(residuals.mean()), float(residuals.std(ddof=1))
    xr = np.linspace(mu-4*sigma, mu+4*sigma, 300)
    bw = (float(residuals.max())-float(residuals.min()))/nb
    ax1.plot(xr, norm.pdf(xr, mu, sigma)*len(residuals)*bw, color=DK, linewidth=2, label="Normal curve")
    ax1.set_title("Histogram of Residuals", fontsize=11, pad=10, color=DK, fontweight="bold")
    ax1.set_xlabel("Residual", fontsize=9); ax1.set_ylabel("Frequency", fontsize=9)
    ax1.tick_params(colors=MT, labelsize=8.5); ax1.legend(fontsize=8)
    for sp in ax1.spines.values(): sp.set_edgecolor("#d0ccc4")
    ax2 = axes[1]; ax2.set_facecolor(AX)
    (osm, osr), (slope, intercept, _) = stats.probplot(residuals, dist="norm")
    ax2.scatter(osm, osr, color=AC, s=28, alpha=0.85, zorder=3, label="Observed")
    xl = np.array([osm[0], osm[-1]])
    ax2.plot(xl, slope*xl+intercept, color=DK, linewidth=1.8, zorder=2, label="Expected")
    ax2.set_title("Normal Q-Q Plot", fontsize=11, pad=10, color=DK, fontweight="bold")
    ax2.set_xlabel("Theoretical Quantiles", fontsize=9); ax2.set_ylabel("Sample Quantiles", fontsize=9)
    ax2.tick_params(colors=MT, labelsize=8.5); ax2.legend(fontsize=8)
    for sp in ax2.spines.values(): sp.set_edgecolor("#d0ccc4")
    plt.tight_layout(pad=2.5)
    return fig

def plot_scatter_residuals(fitted, residuals):
    fig, ax = plt.subplots(figsize=(7.5, 4.2)); style_ax(ax, fig)
    zr = (residuals-residuals.mean())/residuals.std(ddof=1)
    zf = (fitted-fitted.mean())/fitted.std(ddof=1)
    ax.scatter(zf, zr, color=AC, s=30, alpha=0.8, zorder=3, edgecolors="white", linewidths=0.4)
    ax.axhline(0,  color=DK,       linewidth=1.2, linestyle="--", alpha=0.7, label="Zero line")
    ax.axhline(2,  color="#c62828", linewidth=0.8, linestyle=":",  alpha=0.5, label="±2 SD")
    ax.axhline(-2, color="#c62828", linewidth=0.8, linestyle=":",  alpha=0.5)
    ax.set_title("Scatterplot — ZRESID vs ZPRED", fontsize=11, pad=10, color=DK, fontweight="bold")
    ax.set_xlabel("Standardized Predicted Value (ZPRED)", fontsize=9)
    ax.set_ylabel("Standardized Residual (ZRESID)", fontsize=9)
    ax.legend(fontsize=8); plt.tight_layout()
    return fig

def plot_regression_line(x_arr, y_arr, x_label, y_label, model_sm):
    fig, ax = plt.subplots(figsize=(7.5, 4.2)); style_ax(ax, fig)
    ax.scatter(x_arr, y_arr, color=AC, s=35, alpha=0.8, zorder=3,
               edgecolors="white", linewidths=0.4, label="Observed")
    xl   = np.linspace(x_arr.min(), x_arr.max(), 300)
    X_lc = sm.add_constant(xl)
    ax.plot(xl, model_sm.predict(X_lc), color=DK, linewidth=2, zorder=4, label="Regression line")
    ax.set_title(f"Scatter Plot: {x_label} vs {y_label}", fontsize=11, pad=10, color=DK, fontweight="bold")
    ax.set_xlabel(x_label, fontsize=9); ax.set_ylabel(y_label, fontsize=9)
    ax.legend(fontsize=8); plt.tight_layout()
    return fig

def plot_obs_vs_fitted(fitted, y, y_label):
    fig, ax = plt.subplots(figsize=(7.5, 4.2)); style_ax(ax, fig)
    ax.scatter(fitted, y, color=AC, s=30, alpha=0.8, zorder=3,
               edgecolors="white", linewidths=0.4, label="Observations")
    mn = min(float(fitted.min()), float(y.min()))
    mx = max(float(fitted.max()), float(y.max()))
    ax.plot([mn,mx],[mn,mx], color=DK, linewidth=1.8, linestyle="--", label="Perfect fit")
    ax.set_title("Observed vs Fitted Values", fontsize=11, pad=10, color=DK, fontweight="bold")
    ax.set_xlabel("Fitted (Predicted) Values", fontsize=9)
    ax.set_ylabel(f"Observed {y_label}", fontsize=9)
    ax.legend(fontsize=8); plt.tight_layout()
    return fig

def fig_to_bytes(fig):
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor=BG)
    buf.seek(0); return buf.read()


# ═══════════════════════════════════════════════════════════════════════════════
# WORD REPORT
# ═══════════════════════════════════════════════════════════════════════════════

def generate_word_report(reg_type, y_col, x_cols, norm_results, lin_results,
                         vif_df, glejser_df, glejser_passed, dw_value,
                         reg_results, all_passed, figs_bytes):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    def sfmt(v, d=4):
        try:
            f = float(v)
            return "-" if np.isnan(f) else f"{f:.{d}f}"
        except Exception:
            return str(v) if v not in ("",None) else "-"

    def add_h(text, level=1):
        h = doc.add_heading(text, level=level)
        for r in h.runs: r.font.color.rgb = RGBColor(0x1a,0x1a,0x2e)

    def add_p(text, bold=False):
        p = doc.add_paragraph(); r = p.add_run(text); r.bold = bold

    def add_tbl(df, cap=None):
        if cap:
            cp = doc.add_paragraph(cap); cp.runs[0].bold = True
        t = doc.add_table(rows=1, cols=len(df.columns)); t.style = "Table Grid"
        for i,c in enumerate(df.columns):
            t.rows[0].cells[i].text = str(c)
            t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        for _,row in df.iterrows():
            cells = t.add_row().cells
            for i,v in enumerate(row): cells[i].text = sfmt(v)
        doc.add_paragraph()

    title = doc.add_heading(
        f"{'Simple' if reg_type=='Simple' else 'Multiple'} Linear Regression Analysis Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_p(f"Dependent Variable (Y): {y_col}", bold=True)
    add_p(f"Independent Variable(s): {', '.join(x_cols)}", bold=True)
    doc.add_paragraph()

    add_h("1. Normality Test of Residuals")
    n_df = pd.DataFrame([
        {"Test":"Kolmogorov-Smirnov (Lilliefors)","Statistic":norm_results["ks_stat"],"Sig.":norm_results["ks_p"]},
        {"Test":"Shapiro-Wilk","Statistic":norm_results["sw_stat"],"Sig.":norm_results["sw_p"]},
    ])
    add_tbl(n_df, "Tests of Normality")
    pp = norm_results["primary_p"]
    add_p(f"{norm_results['primary']}: Sig. = {pp:.4f} ({'> 0.05' if pp>0.05 else '< 0.05'}). "
          f"Residuals {'are' if norm_results['passed'] else 'are NOT'} normally distributed. "
          f"Assumption: {'SATISFIED' if norm_results['passed'] else 'VIOLATED'}.", bold=True)
    if "normality" in figs_bytes:
        doc.add_picture(BytesIO(figs_bytes["normality"]), width=Inches(5.5))
        doc.add_paragraph()

    add_h("2. Linearity Test")
    for xc, lr in lin_results.items():
        add_h(f"  {xc} → {y_col}", level=2)
        lt = pd.DataFrame([
            {"Source":"Linearity","SS":sfmt(lr["ss_reg"]),"df":lr["df_lin"],"MS":sfmt(lr["ms_lin"]),"F":sfmt(lr["F_lin"]),"Sig.":sfmt(lr["p_lin"])},
            {"Source":"Deviation from Linearity","SS":sfmt(lr.get("ss_dev",np.nan)),"df":lr["df_dev"] if lr["df_dev"]>0 else "-","MS":sfmt(lr.get("ms_dev",np.nan)),"F":sfmt(lr.get("F_dev",np.nan)),"Sig.":sfmt(lr.get("p_dev",np.nan))},
            {"Source":"Within Groups","SS":sfmt(lr.get("ss_within",np.nan)),"df":lr["df_within"],"MS":sfmt(lr.get("ms_within",np.nan)),"F":"-","Sig.":"-"},
        ])
        add_tbl(lt, f"ANOVA Table — {xc} vs {y_col}")
        add_p(f"Sig. Linearity={sfmt(lr['p_lin'])}; Sig. Deviation={sfmt(lr.get('p_dev',np.nan))}. "
              f"Assumption: {'SATISFIED' if lr['passed'] else 'VIOLATED'}.", bold=True)

    if reg_type=="Multiple" and vif_df is not None:
        add_h("3. Multicollinearity Test")
        add_tbl(vif_df.round(4), "Collinearity Statistics")
        mc_ok = all(vif_df["VIF"]<10) and all(vif_df["Tolerance"]>0.10)
        add_p(f"All VIF<10 and Tolerance>0.10: {'Yes' if mc_ok else 'No'}. "
              f"Assumption: {'SATISFIED' if mc_ok else 'VIOLATED'}.", bold=True)

    sec = "4" if reg_type=="Multiple" else "3"
    add_h(f"{sec}. Heteroscedasticity Test (Glejser)")
    add_tbl(glejser_df.round(4), "Glejser Test — Coefficients")
    add_p(f"All predictor Sig.>0.05: {'Yes' if glejser_passed else 'No'}. "
          f"Assumption: {'SATISFIED' if glejser_passed else 'VIOLATED'}.", bold=True)
    if "scatter" in figs_bytes:
        doc.add_picture(BytesIO(figs_bytes["scatter"]), width=Inches(5.0))
        doc.add_paragraph()

    sec2 = "5" if reg_type=="Multiple" else "4"
    add_h(f"{sec2}. Autocorrelation Test (Durbin-Watson)")
    dw_ok = 1.5 <= dw_value <= 2.5
    add_p(f"DW = {dw_value:.4f} ({'within' if dw_ok else 'outside'} 1.5–2.5). "
          f"Assumption: {'SATISFIED' if dw_ok else 'VIOLATED'}.", bold=True)

    sec3 = "6" if reg_type=="Multiple" else "5"
    add_h(f"{sec3}. Regression Analysis Results")
    ms_df = pd.DataFrame([{"R":reg_results["R"],"R Square":reg_results["R2"],
                            "Adjusted R Square":reg_results["adj_R2"],
                            "Std. Error":reg_results["se_est"],"Durbin-Watson":reg_results["dw"]}])
    add_tbl(ms_df.round(4), "Model Summary")
    add_tbl(reg_results["anova"], "ANOVA")
    add_tbl(reg_results["coef_df"], "Coefficients")

    coefs     = reg_results["coef_df"]
    const_val = float(coefs.loc[coefs["Model"]=="(Constant)","B"].values[0])
    eq_parts  = [f"{const_val:.4f}"]
    for _,row in coefs[coefs["Model"]!="(Constant)"].iterrows():
        s = "+" if row["B"]>=0 else ""; eq_parts.append(f"{s}{row['B']:.4f}({row['Model']})")
    eq_str = "Ŷ = " + " ".join(eq_parts)
    add_p(f"Regression Equation: {eq_str}", bold=True)
    if "regression" in figs_bytes:
        doc.add_picture(BytesIO(figs_bytes["regression"]), width=Inches(5.5))
        doc.add_paragraph()

    sec4 = "7" if reg_type=="Multiple" else "6"
    add_h(f"{sec4}. Interpretation and Conclusion")
    sig_model = reg_results["p_F"] < 0.05
    r2_pct    = reg_results["R2"]*100
    lin_s = " ".join(
        f"Linearity {xc}→{y_col}: {'confirmed' if lr['passed'] else 'not confirmed'} (Sig.={sfmt(lr['p_lin'])})."
        for xc,lr in lin_results.items())
    mc_s = ""
    if reg_type=="Multiple" and vif_df is not None:
        mc_s = ("No multicollinearity. " if all(vif_df["VIF"]<10) else "Multicollinearity detected. ")
    pred_s = " ".join(
        f"{row['Model']}: B={row['B']:.4f}, Sig.={float(row['Sig.']):.4f} "
        f"({'significant' if float(row['Sig.'])<0.05 else 'not significant'}, "
        f"{'positive' if row['B']>0 else 'negative'} effect)."
        for _,row in coefs[coefs["Model"]!="(Constant)"].iterrows())
    doc.add_paragraph(
        f"Normality: {norm_results['primary']} Sig.={norm_results['primary_p']:.4f} "
        f"({'satisfied' if norm_results['passed'] else 'violated'}). "
        f"{lin_s} {mc_s}"
        f"Heteroscedasticity: {'satisfied' if glejser_passed else 'violated'}. "
        f"Autocorrelation DW={dw_value:.4f}: {'satisfied' if dw_ok else 'violated'}.\n"
        f"Regression: F={reg_results['F']:.4f}, Sig.={reg_results['p_F']:.4f} "
        f"({'significant' if sig_model else 'not significant'}). "
        f"R²={reg_results['R2']:.4f} ({r2_pct:.1f}% variance explained). "
        f"Equation: {eq_str}. {pred_s}")

    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf.read()


# ═══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def fmt(val, d=4):
    try:
        v = float(val)
        return "-" if np.isnan(v) else f"{v:.{d}f}"
    except Exception:
        return str(val) if val not in ("", None) else "-"

def display_pass_fail(passed):
    tag = "badge-pass" if passed else "badge-fail"
    lbl = "✓ ASSUMPTION SATISFIED" if passed else "✗ ASSUMPTION VIOLATED"
    st.markdown(f'<span class="{tag}">{lbl}</span>', unsafe_allow_html=True)

RECOMMENDATIONS = {
    "normality": "**Recommended actions:**\n1. **Increase sample size** — CLT helps at n > 30–50.\n2. **Transform data** — apply LN or SQRT to Y, re-run all tests.\n3. **Remove outliers** — only if theoretically justified.\n4. **Non-parametric alternative** — Spearman correlation.",
    "linearity": "**Recommended actions:**\n1. **Transform variable** — try LN(X), SQRT(X), or LN(Y).\n2. **Polynomial regression** — add X² term.\n3. **Non-linear regression** — if a specific curve is expected.\n4. **Review theory** — confirm linear relationship is justified.",
    "multicollinearity": "**Recommended actions:**\n1. **Remove one predictor** — drop the less important of correlated predictors.\n2. **Combine variables** — create a composite score.\n3. **Mean-center variables** — subtract column mean from each predictor.\n4. **Increase sample size.**",
    "heteroscedasticity": "**Recommended actions:**\n1. **Transform Y** — LN(Y) or SQRT(Y) to stabilise variance.\n2. **Weighted Least Squares (WLS).**\n3. **Check for omitted variables.**",
    "autocorrelation": "**Recommended actions:**\n1. **Add lagged variable** — include Yt-1 as predictor.\n2. **Difference the data** — transform Y to Yt − Yt-1.\n3. **Time-series models** — ARIMA if data is temporal.\n4. Note: less critical for cross-sectional data.",
}


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    st.markdown("""
    <div style="padding:1.5rem 0 0.5rem 0;">
        <div class="main-title">Regression Analyzer</div>
        <div class="main-subtitle">SPSS-equivalent Simple &amp; Multiple Linear Regression by Muhaimin Abdullah</div>
    </div><hr class="divider">""", unsafe_allow_html=True)

    # Step 1
    st.markdown('<div class="step-badge">Step 01</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-header">Choose Regression Type</div>', unsafe_allow_html=True)
    reg_choice = st.radio("type", ["Simple Regression (1 X, 1 Y)", "Multiple Regression (2+ X, 1 Y)"],
                          horizontal=True, label_visibility="collapsed")
    reg_type = "Simple" if "Simple" in reg_choice else "Multiple"
    st.markdown('<hr class="divider">', unsafe_allow_html=True)

    # Step 2
    st.markdown('<div class="step-badge">Step 02</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-header">Upload CSV Data</div>', unsafe_allow_html=True)
    st.markdown("""
    <div class="interp-box">
    <b>CSV rules:</b> First row = column headers. Each row = one respondent. All analysis columns must be numeric.
    Separator: comma (,). Encoding: UTF-8.<br><br>
    <b>Simple example:</b> <code>respondent_id, Motivation_X, Score_Y</code><br>
    <b>Multiple example:</b> <code>respondent_id, Teaching_X1, Motivation_X2, Score_Y</code>
    </div>""", unsafe_allow_html=True)

    uploaded = st.file_uploader("Upload CSV", type=["csv"], label_visibility="collapsed")
    if uploaded is None:
        st.info("⬆️  Upload a CSV file to begin."); return

    try:
        df = pd.read_csv(uploaded)
    except Exception as e:
        st.error(f"Could not read CSV: {e}"); return

    st.success(f"✓ File loaded — {len(df):,} rows · {len(df.columns)} columns")
    with st.expander("Preview data (first 20 rows)", expanded=False):
        st.dataframe(df.head(20), use_container_width=True)

    numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
    if len(numeric_cols) < 2:
        st.error("Your CSV must contain at least 2 numeric columns."); return

    missing = df[numeric_cols].isnull().sum()
    if missing.any():
        st.warning(f"Missing values detected: {missing[missing>0].to_dict()}. Will be dropped.")
        df = df.dropna(subset=numeric_cols)

    st.markdown('<hr class="divider">', unsafe_allow_html=True)

    # Step 3
    st.markdown('<div class="step-badge">Step 03</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-header">Map Variables</div>', unsafe_allow_html=True)
    c1, c2 = st.columns(2)
    with c1: y_col = st.selectbox("Dependent Variable (Y)", numeric_cols)
    remaining = [c for c in numeric_cols if c != y_col]
    if reg_type == "Simple":
        with c2: x_col = st.selectbox("Independent Variable (X)", remaining)
        x_cols = [x_col]
    else:
        with c2:
            x_cols = st.multiselect("Independent Variables (X1, X2, …)", remaining,
                                    default=remaining[:min(2, len(remaining))])
        if len(x_cols) < 2:
            st.warning("Select at least 2 independent variables."); return
    if not x_cols:
        st.warning("Select at least one independent variable."); return

    y    = df[y_col].astype(float).reset_index(drop=True)
    X_df = df[x_cols].astype(float).reset_index(drop=True)
    st.markdown(f'<div class="interp-box"><b>Model:</b> {y_col} = f({", ".join(x_cols)}) &nbsp;·&nbsp; <b>n = {len(y)}</b></div>',
                unsafe_allow_html=True)
    st.markdown('<hr class="divider">', unsafe_allow_html=True)

    # Step 4 — run analysis
    st.markdown('<div class="step-badge">Step 04 — Automatic</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-header">Prerequisite Assumption Tests</div>', unsafe_allow_html=True)

    figs_bytes = {}; all_passed = True
    with st.spinner("Running full analysis…"):
        reg_results = run_regression(y, X_df)
    residuals = reg_results["residuals"]
    fitted    = reg_results["fitted"]

    # ── Normality ──────────────────────────────────────────────────────────
    st.markdown("### 📋 Test 1 · Normality of Residuals")
    norm_res    = run_normality(residuals)
    norm_passed = norm_res["passed"]
    if not norm_passed: all_passed = False

    ca,cb,cc,cd = st.columns(4)
    ca.metric("n",                         str(norm_res["n"]))
    cb.metric("Shapiro-Wilk W",            f"{norm_res['sw_stat']:.4f}")
    cc.metric("Shapiro-Wilk Sig.",          f"{norm_res['sw_p']:.4f}")
    cd.metric("KS Lilliefors Sig.",         f"{norm_res['ks_p']:.4f}")

    norm_tbl = pd.DataFrame([
        {"Test":"Kolmogorov-Smirnov (Lilliefors)","Statistic":f"{norm_res['ks_stat']:.4f}","df":norm_res['n'],"Sig.":f"{norm_res['ks_p']:.4f}","Recommended when":"n > 50"},
        {"Test":"Shapiro-Wilk",                   "Statistic":f"{norm_res['sw_stat']:.4f}","df":norm_res['n'],"Sig.":f"{norm_res['sw_p']:.4f}","Recommended when":"n ≤ 50"},
    ])
    st.dataframe(norm_tbl, use_container_width=True, hide_index=True)
    display_pass_fail(norm_passed)
    pp = norm_res["primary_p"]
    st.markdown(f"""
    <div class="interp-box">
    <b>Primary test:</b> {norm_res['primary']} (n = {norm_res['n']} {'≤ 50' if norm_res['n']<=50 else '> 50'}).<br>
    Sig. = <b>{pp:.4f}</b> ({'> 0.05 ✓ normally distributed' if pp>0.05 else '< 0.05 ✗ NOT normally distributed'}).<br>
    <b>Note:</b> SPSS applies <i>Lilliefors correction</i> to the KS statistic.
    This app uses <code>statsmodels.stats.diagnostic.lilliefors()</code> — identical correction.
    </div>""", unsafe_allow_html=True)
    if not norm_passed:
        with st.expander("💡 Recommendations — Normality Violated"):
            st.markdown(RECOMMENDATIONS["normality"])

    fig_norm = plot_normality(residuals)
    figs_bytes["normality"] = fig_to_bytes(fig_norm)
    st.pyplot(fig_norm, use_container_width=True); plt.close(fig_norm)
    st.markdown('<hr class="divider">', unsafe_allow_html=True)

    # ── Linearity ──────────────────────────────────────────────────────────
    st.markdown("### 📋 Test 2 · Linearity")
    lin_results = {}
    for xc in x_cols:
        st.markdown(f"**{xc} → {y_col}**")
        lr = run_linearity(X_df[xc], y)
        lin_results[xc] = lr
        if not lr["passed"]: all_passed = False
        ss_comb = lr["ss_reg"] + (lr["ss_dev"] if not np.isnan(lr.get("ss_dev", np.nan)) else 0)
        lin_tbl = pd.DataFrame([
            {"Source":"Between Groups (Combined)","SS":fmt(ss_comb),"df":lr.get("df_groups","-"),"MS":fmt(lr.get("ms_between",np.nan)),"F":fmt(lr.get("F_between",np.nan)),"Sig.":fmt(lr.get("p_between",np.nan))},
            {"Source":"  Linearity","SS":fmt(lr["ss_reg"]),"df":lr["df_lin"],"MS":fmt(lr["ms_lin"]),"F":fmt(lr["F_lin"]),"Sig.":fmt(lr["p_lin"])},
            {"Source":"  Deviation from Linearity","SS":fmt(lr.get("ss_dev",np.nan)),"df":lr["df_dev"] if lr["df_dev"]>0 else "-","MS":fmt(lr.get("ms_dev",np.nan)),"F":fmt(lr.get("F_dev",np.nan)),"Sig.":fmt(lr.get("p_dev",np.nan))},
            {"Source":"Within Groups","SS":fmt(lr.get("ss_within",np.nan)),"df":lr["df_within"],"MS":fmt(lr.get("ms_within",np.nan)),"F":"-","Sig.":"-"},
        ])
        st.dataframe(lin_tbl, use_container_width=True, hide_index=True)
        display_pass_fail(lr["passed"])
        cont_note = ("<br><i>Note: X has all unique values (continuous) — Deviation from Linearity cannot be computed; OLS F-test used instead.</i>"
                     if lr.get("continuous_x") else "")
        p_lin_v = lr["p_lin"]; p_dev_v = lr.get("p_dev", np.nan)
        pls = fmt(p_lin_v); pds = fmt(p_dev_v)
        lin_ok_note = "(< 0.05 ✓)" if (not np.isnan(float(p_lin_v)) and float(p_lin_v)<0.05) else "(> 0.05 ✗)"
        dev_ok_note = "(> 0.05 ✓)" if (pds!="-" and not np.isnan(float(pds)) and float(pds)>0.05) else ""
        st.markdown(f"""
        <div class="interp-box">
        Sig. Linearity = <b>{pls}</b> {lin_ok_note} &nbsp;|&nbsp;
        Sig. Deviation from Linearity = <b>{pds}</b> {dev_ok_note}.{cont_note}<br>
        {xc} → {y_col}: <b>{'linear ✓' if lr['passed'] else 'NOT confirmed as linear ✗'}</b>
        </div>""", unsafe_allow_html=True)
        if not lr["passed"]:
            with st.expander(f"💡 Recommendations — Linearity Violated ({xc})"):
                st.markdown(RECOMMENDATIONS["linearity"])
    st.markdown('<hr class="divider">', unsafe_allow_html=True)

    # ── Multicollinearity ──────────────────────────────────────────────────
    vif_df = None
    if reg_type == "Multiple":
        st.markdown("### 📋 Test 3 · Multicollinearity")
        vif_df    = run_vif(X_df)
        mc_passed = all(vif_df["VIF"]<10) and all(vif_df["Tolerance"]>0.10)
        if not mc_passed: all_passed = False
        vif_disp = vif_df.copy()
        vif_disp["Tolerance"] = vif_disp["Tolerance"].map("{:.4f}".format)
        vif_disp["VIF"]       = vif_disp["VIF"].map("{:.4f}".format)
        st.dataframe(vif_disp, use_container_width=True, hide_index=True)
        display_pass_fail(mc_passed)
        max_vif = float(vif_df["VIF"].max())
        st.markdown(f"""
        <div class="interp-box">
        Criterion: Tolerance > 0.10 AND VIF < 10.<br>
        Max VIF = <b>{max_vif:.4f}</b> ({'< 10 ✓' if max_vif<10 else '≥ 10 ✗'}).
        {'No multicollinearity.' if mc_passed else 'Multicollinearity detected — coefficients unreliable.'}
        </div>""", unsafe_allow_html=True)
        if not mc_passed:
            with st.expander("💡 Recommendations — Multicollinearity"):
                st.markdown(RECOMMENDATIONS["multicollinearity"])
        st.markdown('<hr class="divider">', unsafe_allow_html=True)

    # ── Heteroscedasticity ─────────────────────────────────────────────────
    sec_h = 4 if reg_type=="Multiple" else 3
    st.markdown(f"### 📋 Test {sec_h} · Heteroscedasticity (Glejser + Scatterplot)")
    glejser_df, glejser_passed, _ = run_glejser(residuals, X_df)
    if not glejser_passed: all_passed = False
    g_disp = glejser_df.copy()
    for col in ["B","Std. Error","t","Sig."]: g_disp[col] = g_disp[col].map("{:.4f}".format)
    st.dataframe(g_disp, use_container_width=True, hide_index=True)
    display_pass_fail(glejser_passed)
    sig_vals = glejser_df[glejser_df["Term"]!="(Constant)"]["Sig."].tolist()
    sig_str  = ", ".join(f"{s:.4f}" for s in sig_vals)
    st.markdown(f"""
    <div class="interp-box">
    Predictor Sig. values: <b>{sig_str}</b>
    ({'all > 0.05 ✓ — no heteroscedasticity' if glejser_passed else 'at least one < 0.05 ✗ — heteroscedasticity detected'}).<br>
    Scatterplot (ZRESID vs ZPRED): random scatter around 0 indicates homoscedasticity.
    </div>""", unsafe_allow_html=True)
    if not glejser_passed:
        with st.expander("💡 Recommendations — Heteroscedasticity"):
            st.markdown(RECOMMENDATIONS["heteroscedasticity"])
    fig_scatter = plot_scatter_residuals(fitted, residuals)
    figs_bytes["scatter"] = fig_to_bytes(fig_scatter)
    st.pyplot(fig_scatter, use_container_width=True); plt.close(fig_scatter)
    st.markdown('<hr class="divider">', unsafe_allow_html=True)

    # ── Autocorrelation ────────────────────────────────────────────────────
    sec_a = 5 if reg_type=="Multiple" else 4
    st.markdown(f"### 📋 Test {sec_a} · Autocorrelation (Durbin-Watson)")
    dw_value  = reg_results["dw"]
    dw_passed = 1.5 <= dw_value <= 2.5
    if not dw_passed: all_passed = False
    ca,cb,cc = st.columns(3)
    ca.metric("Durbin-Watson", f"{dw_value:.4f}")
    cb.metric("Acceptable Range", "1.5 – 2.5")
    cc.metric("Status", "PASS ✓" if dw_passed else "FAIL ✗")
    display_pass_fail(dw_passed)
    st.markdown(f"""
    <div class="interp-box">
    DW = <b>{dw_value:.4f}</b> ({'within' if dw_passed else 'outside'} 1.5 – 2.5).<br>
    {'No autocorrelation detected.' if dw_passed else
     ('Positive autocorrelation detected (DW < 1.5).' if dw_value<1.5 else 'Negative autocorrelation detected (DW > 2.5).')}
    </div>""", unsafe_allow_html=True)
    if not dw_passed:
        with st.expander("💡 Recommendations — Autocorrelation"):
            st.markdown(RECOMMENDATIONS["autocorrelation"])
    st.markdown('<hr class="divider">', unsafe_allow_html=True)

    # ── Summary ────────────────────────────────────────────────────────────
    st.markdown("### 📊 Assumption Tests Summary")
    summary = [{"Test":"1 · Normality (Residuals)",
                "Key Statistic":f"{norm_res['primary']} Sig. = {norm_res['primary_p']:.4f}",
                "Result":"✓ SATISFIED" if norm_passed else "✗ VIOLATED"}]
    for xc,lr in lin_results.items():
        summary.append({"Test":f"2 · Linearity ({xc}→{y_col})",
                        "Key Statistic":f"Sig. Linearity = {fmt(lr['p_lin'])}",
                        "Result":"✓ SATISFIED" if lr["passed"] else "✗ VIOLATED"})
    if reg_type=="Multiple" and vif_df is not None:
        mc_ok = all(vif_df["VIF"]<10)
        summary.append({"Test":"3 · Multicollinearity",
                        "Key Statistic":f"Max VIF = {float(vif_df['VIF'].max()):.4f}",
                        "Result":"✓ SATISFIED" if mc_ok else "✗ VIOLATED"})
    summary.append({"Test":f"{sec_h} · Heteroscedasticity (Glejser)",
                    "Key Statistic":f"Min predictor Sig. = {min(sig_vals):.4f}",
                    "Result":"✓ SATISFIED" if glejser_passed else "✗ VIOLATED"})
    summary.append({"Test":f"{sec_a} · Autocorrelation (DW)",
                    "Key Statistic":f"DW = {dw_value:.4f}",
                    "Result":"✓ SATISFIED" if dw_passed else "✗ VIOLATED"})
    st.dataframe(pd.DataFrame(summary), use_container_width=True, hide_index=True)
    if all_passed:
        st.success("✅ All assumptions satisfied — regression results below are fully valid.")
    else:
        st.warning("⚠️ One or more assumptions violated. See recommendations above. Results shown for reference.")
    st.markdown('<hr class="divider">', unsafe_allow_html=True)

    # ── Regression Results ─────────────────────────────────────────────────
    sec_r = 6 if reg_type=="Multiple" else 5
    st.markdown(f'<div class="section-header">{sec_r} · Regression Analysis Results</div>', unsafe_allow_html=True)

    st.markdown("#### Model Summary")
    mc1,mc2,mc3,mc4,mc5 = st.columns(5)
    mc1.metric("R",                  f"{reg_results['R']:.4f}")
    mc2.metric("R Square",           f"{reg_results['R2']:.4f}")
    mc3.metric("Adjusted R²",        f"{reg_results['adj_R2']:.4f}")
    mc4.metric("Std. Error of Est.", f"{reg_results['se_est']:.4f}")
    mc5.metric("Durbin-Watson",      f"{reg_results['dw']:.4f}")
    r2_pct = reg_results["R2"]*100
    st.markdown(f"""
    <div class="interp-box">
    <b>R = {reg_results['R']:.4f}</b> — correlation strength between predicted and actual {y_col}.<br>
    <b>R² = {reg_results['R2']:.4f}</b> — <b>{r2_pct:.1f}%</b> of variance in <b>{y_col}</b> explained by {', '.join(x_cols)}.<br>
    <b>Adjusted R² = {reg_results['adj_R2']:.4f}</b> — corrected for number of predictors and sample size.
    </div>""", unsafe_allow_html=True)
    st.markdown('<hr class="divider">', unsafe_allow_html=True)

    st.markdown("#### ANOVA")
    anova_disp = reg_results["anova"].copy()
    for col in ["Sum of Squares","Mean Square","F","Sig."]:
        anova_disp[col] = anova_disp[col].apply(lambda v: fmt(v) if v!="" else "-")
    anova_disp["df"] = anova_disp["df"].apply(lambda v: str(int(v)) if v!="" else "-")
    st.dataframe(anova_disp, use_container_width=True, hide_index=True)
    sig_model = reg_results["p_F"] < 0.05
    st.markdown(f"""
    <div class="interp-box">
    F = <b>{reg_results['F']:.4f}</b>, Sig. = <b>{reg_results['p_F']:.4f}</b>
    ({'< 0.05' if sig_model else '> 0.05'}).<br>
    Model is <b>{'statistically significant ✓' if sig_model else 'NOT statistically significant ✗'}</b>.
    </div>""", unsafe_allow_html=True)
    st.markdown('<hr class="divider">', unsafe_allow_html=True)

    st.markdown("#### Coefficients")
    coefs = reg_results["coef_df"]
    coef_disp = coefs.copy()
    for col in ["B","Std. Error","t"]: coef_disp[col] = coef_disp[col].map("{:.4f}".format)
    coef_disp["Beta"]  = coef_disp["Beta"].apply(lambda v: f"{v:.4f}" if v!="" else "-")
    coef_disp["Sig."]  = coef_disp["Sig."].map(lambda v: f"{v:.4f} *" if float(v)<0.05 else f"{v:.4f}")
    st.dataframe(coef_disp, use_container_width=True, hide_index=True)
    st.caption("* Sig. < 0.05 — statistically significant")

    const_val = float(coefs.loc[coefs["Model"]=="(Constant)","B"].values[0])
    eq_parts  = [f"{const_val:.4f}"]
    interp_parts = []
    for _,row in coefs[coefs["Model"]!="(Constant)"].iterrows():
        s = "+" if row["B"]>=0 else ""
        eq_parts.append(f"{s}{row['B']:.4f}({row['Model']})")
        sig_x = float(row["Sig."]) < 0.05
        interp_parts.append(
            f"<b>{row['Model']}</b>: B={row['B']:.4f}, β={fmt(row['Beta'])}, "
            f"Sig.={float(row['Sig.']):.4f} → <b>{'significant' if sig_x else 'not significant'}</b>, "
            f"{'positive' if row['B']>0 else 'negative'} effect on {y_col}.")
    eq_str = "Ŷ = " + " ".join(eq_parts)
    st.markdown(f'<div class="equation-box">{eq_str}</div>', unsafe_allow_html=True)
    st.markdown(f"""
    <div class="interp-box">
    <b>Constant (a) = {const_val:.4f}:</b> predicted {y_col} when all X = 0.<br><br>
    {"<br>".join(interp_parts)}
    </div>""", unsafe_allow_html=True)
    st.markdown('<hr class="divider">', unsafe_allow_html=True)

    # Plot
    if reg_type == "Simple":
        x_plot = X_df[x_cols[0]].values; y_plot = y.values
        m_plot = sm.OLS(y_plot, sm.add_constant(x_plot)).fit()
        fig_reg = plot_regression_line(x_plot, y_plot, x_cols[0], y_col, m_plot)
    else:
        fig_reg = plot_obs_vs_fitted(fitted, y, y_col)
    figs_bytes["regression"] = fig_to_bytes(fig_reg)
    st.pyplot(fig_reg, use_container_width=True); plt.close(fig_reg)
    st.markdown('<hr class="divider">', unsafe_allow_html=True)

    # ── Conclusion ─────────────────────────────────────────────────────────
    sec_c = 7 if reg_type=="Multiple" else 6
    st.markdown(f'<div class="section-header">{sec_c} · Interpretation & Conclusion</div>', unsafe_allow_html=True)
    lin_snts = " ".join(
        f"Linearity {xc}→{y_col}: {'confirmed' if lr['passed'] else 'not confirmed'} (Sig. Linearity={fmt(lr['p_lin'])}; Sig. Deviation={fmt(lr.get('p_dev',np.nan))})."
        for xc,lr in lin_results.items())
    mc_snt = ""
    if reg_type=="Multiple" and vif_df is not None:
        mc_snt = ("No multicollinearity (all VIF < 10). " if all(vif_df["VIF"]<10)
                  else "Multicollinearity detected. ")
    pred_snts = " ".join(
        f"{row['Model']}: B={row['B']:.4f}, β={fmt(row['Beta'])}, t={row['t']:.4f}, Sig.={float(row['Sig.']):.4f} "
        f"→ {'significant' if float(row['Sig.'])<0.05 else 'not significant'}, "
        f"{'positive' if row['B']>0 else 'negative'} effect on {y_col}."
        for _,row in coefs[coefs["Model"]!="(Constant)"].iterrows())
    conclusion = (
        f"Prior to the {'simple' if reg_type=='Simple' else 'multiple'} linear regression analysis, "
        f"all prerequisite assumption tests were performed. "
        f"The {norm_res['primary']} normality test: Sig. = {norm_res['primary_p']:.4f} "
        f"({'satisfied' if norm_passed else 'violated'}). "
        f"{lin_snts} {mc_snt}"
        f"Glejser test: {'no heteroscedasticity' if glejser_passed else 'heteroscedasticity present'}. "
        f"Durbin-Watson = {dw_value:.4f}: {'no autocorrelation' if dw_passed else 'autocorrelation detected'}.\n\n"
        f"Regression: F = {reg_results['F']:.4f}, Sig. = {reg_results['p_F']:.4f} "
        f"({'significant' if sig_model else 'not significant'}). "
        f"R² = {reg_results['R2']:.4f} ({r2_pct:.1f}% variance in {y_col} explained). "
        f"Equation: {eq_str}. {pred_snts}")
    st.markdown(f'<div class="interp-box" style="font-size:0.94rem;line-height:1.8;">'
                f'{conclusion.replace(chr(10),"<br><br>")}</div>', unsafe_allow_html=True)
    st.markdown('<hr class="divider">', unsafe_allow_html=True)

    # ── Downloads ──────────────────────────────────────────────────────────
    st.markdown('<div class="section-header">⬇ Download Report</div>', unsafe_allow_html=True)
    with st.spinner("Generating Word report…"):
        word_bytes = generate_word_report(
            reg_type=reg_type, y_col=y_col, x_cols=x_cols,
            norm_results=norm_res, lin_results=lin_results,
            vif_df=vif_df, glejser_df=glejser_df,
            glejser_passed=glejser_passed, dw_value=dw_value,
            reg_results=reg_results, all_passed=all_passed,
            figs_bytes=figs_bytes)

    dl1, dl2 = st.columns(2)
    with dl1:
        st.download_button("⬇  Word Report (.docx)", data=word_bytes,
            file_name=f"regression_report_{y_col}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True)
    with dl2:
        rows = []
        for k,v in [("R",reg_results["R"]),("R Square",reg_results["R2"]),
                    ("Adjusted R Square",reg_results["adj_R2"]),
                    ("Std. Error",reg_results["se_est"]),("Durbin-Watson",reg_results["dw"])]:
            rows.append({"Section":"MODEL SUMMARY","Item":k,"Value":f"{v:.4f}"})
        rows += [{"Section":"ANOVA","Item":"F","Value":f"{reg_results['F']:.4f}"},
                 {"Section":"ANOVA","Item":"Sig.","Value":f"{reg_results['p_F']:.4f}"}]
        for _,row in coefs.iterrows():
            rows.append({"Section":"COEFFICIENTS","Item":f"{row['Model']} B","Value":f"{row['B']:.4f}"})
            rows.append({"Section":"COEFFICIENTS","Item":f"{row['Model']} Sig.","Value":f"{float(row['Sig.']):.4f}"})
        csv_bytes = pd.DataFrame(rows).to_csv(index=False).encode("utf-8")
        st.download_button("⬇  Results Summary (.csv)", data=csv_bytes,
            file_name=f"regression_results_{y_col}.csv", mime="text/csv",
            use_container_width=True)

    st.markdown('<div style="text-align:center;color:#aaa;font-size:0.76rem;padding:2rem 0 1rem 0;font-family:monospace;">Regression Analyzer · SPSS-equivalent · statsmodels + scipy</div>',
                unsafe_allow_html=True)


if __name__ == "__main__":
    main()
